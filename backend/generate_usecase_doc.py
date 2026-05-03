from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles helper ──
def heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_table(headers, data, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
            for para in row[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)
    doc.add_paragraph()

# ════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("DataLens")
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("AI-Powered Survey Insights Platform")
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_paragraph()
tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = tagline.add_run("Use Case & Product Overview — All Modules")
run3.font.size = Pt(13)
run3.font.italic = True
run3.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = date_para.add_run(f"BorderlessAccess  |  {datetime.date.today().strftime('%B %Y')}")
run4.font.size = Pt(11)
run4.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

doc.add_page_break()

# ════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════
heading("1. Executive Summary")
para(
    "DataLens is an AI-powered web platform that transforms raw survey data into instant, "
    "actionable insights. Designed for BorderlessAccess researchers and clients, DataLens "
    "eliminates the dependency on SPSS expertise by allowing users to ask natural language "
    "questions and receive data-backed answers, charts, and summaries in seconds."
)
para(
    "The platform is delivered in three progressive modules, each adding a distinct capability "
    "layer — from conversational AI querying (Module 1) to automated reporting (Module 2) to "
    "self-serve crosstab building (Module 3). Modules are built and released sequentially, "
    "incorporating real-user feedback before the next is developed."
)

doc.add_paragraph()
heading("Key Business Benefits", level=2)
bullet("Dramatically reduces time-to-insight from days to seconds")
bullet("Removes SPSS expertise barrier for clients and junior researchers")
bullet("Enables clients to self-serve, reducing analyst dependency")
bullet("Supports weighted analysis and wave-over-wave trends out of the box")
bullet("Role-based access ensures data security and client data isolation")

doc.add_page_break()

# ════════════════════════════════════════════
# 2. PLATFORM OVERVIEW
# ════════════════════════════════════════════
heading("2. Platform Overview")
para(
    "DataLens ingests SPSS (.sav) survey data along with a structured datamap that describes "
    "each variable, its question text, answer options, and type. Once uploaded, the data is "
    "immediately available for querying via the AI chatbot, automated summary, or the crosstab builder."
)

doc.add_paragraph()
heading("Data Input", level=2)
bullet("Survey data: SPSS .sav file (single or merged multi-wave file)")
bullet("Datamap: Excel file describing variables, question text, answer options, and weight flag")
bullet("Wave studies: a single merged SPSS file with a wave variable identifying each respondent's wave")
bullet("Weighting: weight variable lives inside the SPSS file, flagged in the datamap")

doc.add_paragraph()
heading("User Roles (Phase B — Access Control)", level=2)
add_table(
    ["Role", "Capabilities"],
    [
        ["Super Admin", "Full system access — all users, all projects, system settings"],
        ["Supervisor", "Create projects, upload data, assign Viewers, set data expiry, view all chat history"],
        ["Viewer", "Access assigned projects only, run queries, view own chat history"],
    ],
    col_widths=[1.5, 5.0]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 3. MODULE 1 — AI CHATBOT
# ════════════════════════════════════════════
heading("3. Module 1 — AI Chatbot for Survey Insights")
heading("Overview", level=2)
para(
    "Module 1 is the core of DataLens. Researchers and clients ask natural language questions "
    "about survey data and receive instant, accurate answers — complete with a narrative summary, "
    "base statement, and visualisation. No SPSS knowledge required."
)

doc.add_paragraph()
heading("How It Works", level=2)
bullet("User types a question in plain English (e.g. \"What % of people from UAE heard of AlUla?\")")
bullet("Claude AI identifies the relevant SPSS variables using the uploaded datamap")
bullet("Filters are applied to the dataset based on the question conditions")
bullet("The appropriate statistical function is run on the filtered data")
bullet("Claude returns a narrative text answer, base statement, and chart JSON")
bullet("A bar, pie, or line chart is rendered instantly in the browser")

doc.add_paragraph()
heading("Supported Query Types", level=2)
add_table(
    ["Query Type", "Example", "Output"],
    [
        ["Frequency / %", "What % of respondents have visited Asia?", "Bar or pie chart with %"],
        ["Filtered frequency", "Among UAE respondents, which destinations are most considered?", "Filtered bar chart"],
        ["Top-box score", "What is the top-2 box score for ad likeability?", "% with base N"],
        ["Mean score", "What is the average rating for ad clarity?", "Mean score with base N"],
        ["Wave-over-wave trend", "How has awareness of AlUla changed across waves?", "Line chart by wave"],
    ],
    col_widths=[1.8, 2.8, 1.8]
)

doc.add_paragraph()
heading("Sample Questions (from Demo Dataset)", level=2)
bullet("Which leisure destinations within Saudi Arabia are most widely heard of?")
bullet("What % of respondents who travelled internationally are likely to visit again in the next 12 months?")
bullet("How does awareness of Neom City differ between KSA and UAE respondents?")
bullet("What is the top-2 box score for how interesting the advertisement was?")
bullet("Among respondents with a postgraduate education, which destinations have they visited in the past 3 years?")

doc.add_paragraph()
heading("Key Design Decisions", level=2)
bullet("Base = respondents satisfying the filter, not total sample — always stated explicitly in the response")
bullet("Multi-response questions: base = total filtered respondents; % can sum to more than 100%")
bullet("Both weighted % and unweighted base N shown in every response")
bullet("Chat history persisted per project per user for audit and continuity")

doc.add_paragraph()
heading("Build Status", level=2)
bullet("Phase A (Core Engine): Complete — backend API, stats engine, Claude AI agent all built")
bullet("Phase B (Access Controls): Pending — JWT auth, role-based middleware, Ops admin panel")
bullet("Frontend: In progress — React chat UI with chart panel")

doc.add_page_break()

# ════════════════════════════════════════════
# 4. MODULE 2 — AUTOMATED SURVEY SUMMARY
# ════════════════════════════════════════════
heading("4. Module 2 — Automated Full Survey Summary")
heading("Overview", level=2)
para(
    "Module 2 auto-generates a comprehensive insights report across all questions in the survey. "
    "Rather than asking individual questions, the researcher or client receives a full summary "
    "of the dataset — covering all key findings, patterns, and notable results — in one action."
)

doc.add_paragraph()
heading("How It Works", level=2)
bullet("Triggered after data upload (or manually by the researcher)")
bullet("Claude AI iterates through all questions in the datamap and generates insights for each")
bullet("Global filter option: rerun the entire summary on a specific sub-group (e.g. females only, UAE only)")
bullet("Output displayed on-screen as a structured dashboard or exported as a report")

doc.add_paragraph()
heading("Key Use Cases", level=2)
bullet("Client debrief preparation — generate a ready-to-present summary without manual analysis")
bullet("Quick dataset review — understand the shape of new data immediately after upload")
bullet("Filtered sub-group reporting — rerun summary on any demographic or behavioural segment")

doc.add_paragraph()
heading("Access Controls", level=2)
bullet("Same role-based access as Module 1 — Viewers see only their assigned projects")
bullet("Supervisors can download or share the summary report")

doc.add_paragraph()
heading("Build Status", level=2)
bullet("Pending — will be scoped and built after Module 1 is live and user feedback is collected")

doc.add_page_break()

# ════════════════════════════════════════════
# 5. MODULE 3 — CROSSTAB BUILDER
# ════════════════════════════════════════════
heading("5. Module 3 — Self-Serve Crosstab & Chart Builder")
heading("Overview", level=2)
para(
    "Module 3 gives researchers and advanced clients a drag-and-drop interface to build their "
    "own crosstabs, apply filters and weights, and generate charts — without writing a single "
    "line of code or opening SPSS. This is the most powerful and flexible module, designed for "
    "users who need full analytical control."
)

doc.add_paragraph()
heading("How It Works", level=2)
bullet("User selects questions from a panel and drags them onto rows and columns")
bullet("Filters and weights are applied independently per crosstab")
bullet("System generates the crosstab table with weighted % and base N")
bullet("Charts are generated from the crosstab in one click")
bullet("Results can be exported to Excel (standard or banner table format)")

doc.add_paragraph()
heading("Key Capabilities", level=2)
add_table(
    ["Capability", "Description"],
    [
        ["Single-code questions", "Standard row × column crosstab with weighted %"],
        ["Multi-response questions", "Base = total respondents; % can exceed 100%"],
        ["Grid / matrix questions", "Rows = items, columns = scale points; row-level analysis supported"],
        ["Weight application", "Apply or remove the weight variable per analysis"],
        ["Custom filters", "Filter by any variable before building the crosstab"],
        ["Chart generation", "Bar, stacked bar, pie, or line chart from any crosstab"],
        ["Excel export", "Standard table or formatted banner table"],
    ],
    col_widths=[2.2, 4.4]
)

doc.add_paragraph()
heading("Build Status", level=2)
bullet("Pending — the most complex module, built last after Module 1 and 2 feedback")
bullet("Open design questions (significance testing, grid row analysis) to be confirmed during scoping")

doc.add_page_break()

# ════════════════════════════════════════════
# 6. TECH STACK
# ════════════════════════════════════════════
heading("6. Technology Stack")
add_table(
    ["Component", "Technology", "Purpose"],
    [
        ["Backend API", "FastAPI (Python)", "REST API, data processing, business logic"],
        ["AI Engine", "Claude claude-sonnet-4-6 (Anthropic)", "Natural language understanding, tool use, response generation"],
        ["Frontend", "React + TypeScript", "Chat interface, chart rendering, file upload"],
        ["Charts", "Recharts", "Bar, pie, line chart visualisations"],
        ["Database", "SQL Server (SQLAlchemy)", "Projects, users, access grants, chat history"],
        ["Data Processing", "pyreadstat + pandas", "SPSS file parsing, weighted statistics"],
        ["Deployment", "On-premise server", "Internal hosting, no cloud dependency"],
    ],
    col_widths=[1.8, 2.2, 2.6]
)

doc.add_page_break()

# ════════════════════════════════════════════
# 7. DATA MANAGEMENT NOTE
# ════════════════════════════════════════════
heading("7. Data Management — Re-upload & Revision Support")
para(
    "An important operational requirement for DataLens is the ability to re-upload SPSS data "
    "for an existing project without losing any project configuration or chat history. "
    "This is a common scenario in market research and must be supported from Day 1."
)

doc.add_paragraph()
heading("Re-upload Scenarios", level=2)
bullet("Quality control: records removed after identifying invalid or duplicate responses")
bullet("Sample augmentation: new respondents added by the client after initial fieldwork")
bullet("Response revision: individual answers corrected after performing sanity or consistency checks")
bullet("Variable changes: new derived variables or recoded variables added to the dataset")

doc.add_paragraph()
heading("How Re-upload Works", level=2)
bullet("Supervisor or Ops uploads a new SPSS file and datamap to the existing project")
bullet("The system replaces the processed data (parquet file) and refreshes the Question Registry")
bullet("All existing project settings, user assignments, and access controls remain unchanged")
bullet("Chat history is preserved — prior queries are not deleted")
bullet("A timestamp and upload log records each data version for audit purposes")

doc.add_paragraph()
para(
    "Note: If variable names or question structures change significantly in the revised SPSS file, "
    "the datamap must also be updated and re-uploaded alongside the new SPSS file.",
    italic=True
)

doc.add_page_break()

# ════════════════════════════════════════════
# 8. ROLLOUT PLAN WITH TIMELINES
# ════════════════════════════════════════════
heading("8. Rollout Plan & Timeline")
para(
    "Timelines below are estimates based on standard development effort. "
    "Actual timelines depend on the resource model chosen (see Section 9). "
    "All estimates assume focused, uninterrupted development effort."
)
doc.add_paragraph()
add_table(
    ["Phase", "Scope", "Est. Duration", "Cumulative Week"],
    [
        ["POC / Demo",
         "Sample data upload, AI chatbot query, chart output. No auth. For management sign-off.",
         "1–2 weeks", "Week 2"],
        ["Module 1 — Phase A",
         "Full chatbot with real project data. All query types (frequency, top-box, mean, trend). Data re-upload support.",
         "4–6 weeks", "Week 8"],
        ["Module 1 — Phase B",
         "JWT authentication, 3 user roles, Ops admin panel, data expiry & access revocation.",
         "3–4 weeks", "Week 12"],
        ["UAT & Stabilisation",
         "User acceptance testing with internal team and pilot clients. Bug fixes and refinements.",
         "2 weeks", "Week 14"],
        ["Module 2",
         "Automated full survey summary. Global filter to rerun on sub-groups. Report export.",
         "4–6 weeks", "Week 20"],
        ["Module 3",
         "Self-serve crosstab & chart builder. Multi-response, grid, weight support. Excel export.",
         "8–12 weeks", "Week 32"],
    ],
    col_widths=[1.6, 3.2, 1.2, 1.4]
)

doc.add_paragraph()
para(
    "Total estimated duration: 28–32 weeks (~7–8 months) from POC to full platform delivery. "
    "Each module is released and validated with real users before the next is built.",
    italic=True
)

doc.add_page_break()

# ════════════════════════════════════════════
# 9. RESOURCE PLANNING
# ════════════════════════════════════════════
heading("9. Resource Planning")
para(
    "Two resource models are presented below. The appropriate model depends on the organisation's "
    "decision to use an external development team or leverage internal resources with domain expertise."
)

doc.add_paragraph()
heading("Scenario A — Pure Development Team (External / Dedicated)", level=2)
para(
    "A dedicated development team is engaged for the full build. The internal team's role is "
    "limited to requirements sign-off, UAT, and providing domain knowledge on survey variables and business rules."
)
doc.add_paragraph()
add_table(
    ["Role", "Profile & Skills", "M1-A", "M1-B", "M2", "M3"],
    [
        ["Senior Full-stack Developer",
         "5+ yrs Python (FastAPI) + React/TypeScript. REST APIs, SQL Server, data pipelines.",
         "5 wks", "3 wks", "4 wks", "10 wks"],
        ["AI / Prompt Engineer",
         "Claude / LLM API integration, tool-use patterns, prompt design, streaming responses.",
         "4 wks", "1 wk", "4 wks", "2 wks"],
        ["QA Engineer (part-time)",
         "API testing (Postman/pytest), UI testing, survey domain awareness for test case design.",
         "2 wks", "2 wks", "2 wks", "4 wks"],
        ["BA / Project Coordinator (part-time)",
         "Requirements documentation, sprint planning, UAT coordination, stakeholder communication.",
         "Throughout", "Throughout", "Throughout", "Throughout"],
    ],
    col_widths=[1.9, 3.0, 0.7, 0.7, 0.6, 0.6]
)

doc.add_paragraph()
heading("Scenario B — Internal Team (Domain Experts, No Development Background)", level=2)
para(
    "The internal team — who deeply understand the survey data, variables, and business context — "
    "supports the build by owning requirements, testing, and domain validation. "
    "A single contracted developer is engaged for the code. Timeline is approximately 30–40% longer "
    "than Scenario A due to knowledge transfer overhead, but total cost is significantly lower."
)
doc.add_paragraph()
add_table(
    ["Role", "Profile", "Contribution to Project"],
    [
        ["Research Lead / Data Analyst",
         "Experienced in survey research, SPSS data, crosstabs",
         "Defines question mappings, validates AI outputs, writes test cases based on known data"],
        ["Survey Operations",
         "Handles SPSS file preparation, fieldwork data",
         "Prepares SPSS files and datamaps, validates variable structures, tests re-upload scenarios"],
        ["Product Owner / Manager",
         "Understands client and researcher workflows",
         "Reviews features at each sprint, signs off UAT, sets module priorities"],
        ["Contracted Developer (1 person)",
         "Full-stack developer with Python + React experience",
         "Sole developer; internal team provides requirements, test data, and domain review"],
    ],
    col_widths=[1.7, 2.0, 2.9]
)

doc.add_paragraph()
para(
    "Note: Scenario B works well for Modules 1 and 2. For Module 3 (Crosstab Builder), "
    "the higher complexity may warrant adding a second developer or extending the timeline.",
    italic=True
)

doc.add_page_break()

# ════════════════════════════════════════════
# 10. COST ESTIMATION
# ════════════════════════════════════════════
heading("10. Cost Estimation")
para(
    "All figures below are indicative estimates. Actual costs will vary based on geography, "
    "vendor rates, and final scope. Costs are split into four categories: development team, "
    "AI API usage, infrastructure, and tools & licensing."
)

doc.add_paragraph()
heading("A. Development Team Cost", level=2)
para("Applicable only under Scenario A (pure development team). Under Scenario B, replace with contracted developer rate.")
doc.add_paragraph()
add_table(
    ["Role", "Est. Monthly Rate", "Engaged For", "Est. Total Cost"],
    [
        ["Senior Full-stack Developer", "$3,500 – $6,000 / month", "~6 months (M1+M2+M3)", "$21,000 – $36,000"],
        ["AI / Prompt Engineer", "$3,500 – $6,000 / month", "~4 months (M1+M2)", "$14,000 – $24,000"],
        ["QA Engineer (50% time)", "$1,000 – $2,000 / month", "~6 months", "$6,000 – $12,000"],
        ["BA / Project Coordinator (50%)", "$1,000 – $1,800 / month", "~7 months", "$7,000 – $12,600"],
        ["Total (Scenario A)", "", "Full Platform", "$48,000 – $84,600"],
    ],
    col_widths=[2.2, 1.8, 1.8, 1.8]
)

doc.add_paragraph()
heading("B. Anthropic API Usage Cost (Per Query)", level=2)
para(
    "DataLens uses Claude claude-sonnet-4-6 via the Anthropic API. Cost is pay-per-use based on tokens processed. "
    "Estimates below assume typical survey question complexity."
)
doc.add_paragraph()
add_table(
    ["Module", "API Calls Per Action", "Est. Cost Per Action", "Est. Monthly Cost (100 actions/day)"],
    [
        ["Module 1 — Chatbot query",
         "1–2 calls per user question",
         "$0.01 – $0.05 per query",
         "$30 – $150 / month"],
        ["Module 2 — Full survey summary",
         "50–100 calls per summary run (1 per question)",
         "$0.50 – $2.00 per summary",
         "$10 – $40 / month (if run ~10×/month)"],
        ["Module 3 — Crosstab Builder",
         "No AI API calls (pure computation)",
         "$0.00",
         "$0.00"],
    ],
    col_widths=[1.8, 1.8, 1.6, 2.4]
)
doc.add_paragraph()
para(
    "Note: API costs scale with usage volume. For a team of 10–20 active users, "
    "total monthly API spend is estimated at $50–$200/month across all modules.",
    italic=True
)

doc.add_paragraph()
heading("C. Infrastructure & Server Cost", level=2)
add_table(
    ["Item", "Type", "Est. Cost"],
    [
        ["On-premise application server", "One-time hardware (if new server needed)", "$2,000 – $5,000"],
        ["OR Cloud VM (e.g. Azure/AWS)", "Monthly recurring (if cloud hosted)", "$150 – $400 / month"],
        ["SQL Server", "Already licensed — no additional cost", "$0"],
        ["Domain name + SSL certificate", "Annual (if externally accessible)", "$50 – $100 / year"],
    ],
    col_widths=[2.4, 2.4, 1.8]
)

doc.add_paragraph()
heading("D. Tools & Licensing Cost", level=2)
add_table(
    ["Tool / Library", "Cost", "Notes"],
    [
        ["Python, FastAPI, React, Recharts", "Free (open source)", "All core development frameworks are open source"],
        ["Anthropic API", "Pay-per-use", "Covered in Section B above"],
        ["pyreadstat (SPSS parsing)", "Free (open source)", "No SPSS licence required on the server"],
        ["SQL Server", "Already licensed", "Using existing BorderlessAccess SQL Server instance"],
        ["VS Code / development IDE", "Free", "Standard developer tooling"],
        ["Git / version control", "Free", "GitHub free tier or internal Git server"],
    ],
    col_widths=[2.4, 1.4, 2.8]
)

doc.add_paragraph()
heading("Cost Summary", level=2)
add_table(
    ["Category", "Scenario A (Dev Team)", "Scenario B (Internal + 1 Dev)"],
    [
        ["Development Team", "$48,000 – $84,600 (full platform)", "$15,000 – $30,000 (contracted dev only)"],
        ["Anthropic API (monthly, ongoing)", "$50 – $200 / month", "$50 – $200 / month"],
        ["Infrastructure (one-time or monthly)", "$2,000–$5,000 or $150–$400/mo", "Same"],
        ["Tools & Licensing", "Minimal — mostly open source", "Same"],
        ["Estimated Total (first year)", "$60,000 – $100,000", "$25,000 – $45,000"],
    ],
    col_widths=[2.2, 2.4, 2.4]
)

doc.add_paragraph()
para(
    "All cost figures are indicative and should be validated with actual vendor quotes. "
    "Scenario B is the recommended starting point for BorderlessAccess given the strong "
    "internal domain expertise available in the research and operations teams.",
    italic=True
)

# ════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════
out_path = r"C:\Users\KarthikShanmugam\ClaudePOC\DataLens\Module 1\DataLens_UseCase_Document.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
