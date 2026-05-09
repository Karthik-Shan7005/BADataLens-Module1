"""
Generate DataLens Module 1 — Business Case & Roadmap (Word .docx)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT = r"C:\Users\KarthikShanmugam\ClaudePOC\DataLens\Module 1\DataLens_Module1_BusinessCase.docx"

# ── colour palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1a, 0x1f, 0x3a)   # sidebar dark navy
TEAL   = RGBColor(0x00, 0x8b, 0x8b)   # accent
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xF0, 0xF4, 0xFF)   # table alt row
GREY   = RGBColor(0x64, 0x74, 0x8B)
GREEN  = RGBColor(0x16, 0xa3, 0x4a)
RED    = RGBColor(0xdc, 0x26, 0x26)
AMBER  = RGBColor(0xd9, 0x77, 0x06)

doc = Document()

# ── page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(text, level=1, color=NAVY):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def body(text, size=10.5, color=None, bold=False, italic=False, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, size=10.5):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.size = Pt(size)
    return p

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1a1f3a')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('DataLens')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('AI Survey Insights Platform')
run2.font.size = Pt(18)
run2.font.color.rgb = TEAL

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('Module 1 — Business Case & Investment Proposal')
run3.font.size = Pt(14)
run3.font.bold = True
run3.font.color.rgb = GREY

doc.add_paragraph()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('Prepared by: Borderless Access — Panel Analytics Team')
run4.font.size = Pt(11)
run4.font.color.rgb = GREY

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p5.add_run('May 2026  |  Confidential')
run5.font.size = Pt(10)
run5.font.color.rgb = GREY
run5.font.italic = True

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# 1. WHAT IS DATALENS MODULE 1?
# ═══════════════════════════════════════════════════════════════════════════
heading('1. What Is DataLens Module 1?', level=1)
divider()
body(
    'DataLens Module 1 is an AI-powered survey analytics chatbot that allows researchers and '
    'clients to ask natural language questions about SPSS survey data and receive instant, '
    'weighted statistical insights — with supporting charts — through a browser-based interface.',
    size=11
)
doc.add_paragraph()
body('Key capabilities delivered in Module 1:', bold=True, size=10.5)
bullet('Natural language querying — ask questions in plain English, get data-backed answers')
bullet('Weighted frequency tables — Response | Count (n) | Percentage (%) with rim-weight support')
bullet('Top Box / Bottom Box scoring — e.g. Top 2 Box on a 5-point satisfaction scale')
bullet('Mean / average scores for numeric and rating questions')
bullet('Wave-over-wave trend analysis for multi-wave tracker studies')
bullet('Dual base reporting — total respondents vs. answered respondents (skip-logic aware)')
bullet('Auto-generated bar and line charts alongside every response')
bullet('Full SPSS (.sav) + datamap ingestion — no manual data re-entry')
bullet('Multi-project support — multiple studies, each with independent data and access')
bullet('Query history — all questions and answers saved per project')

doc.add_paragraph()
body(
    'The platform is built on Anthropic\'s Claude Sonnet AI model with a FastAPI (Python) '
    'backend and a React + TypeScript frontend — a modern, maintainable technology stack.',
    size=10.5, color=GREY, italic=True
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# 2. PROBLEM & SOLUTION
# ═══════════════════════════════════════════════════════════════════════════
heading('2. Current Challenges & How DataLens Solves Them', level=1)
divider()
body(
    'Market research analysts and their clients face a consistent set of pain points when working '
    'with survey data today. These challenges slow down decision-making, increase analyst workload, '
    'and limit clients\' ability to explore data independently.',
    size=10.5
)
doc.add_paragraph()

challenges = [
    (
        'Slow, manual data analysis',
        'Generating a weighted frequency table or cross-tab in SPSS or Excel requires syntax knowledge, '
        'manual steps, and significant time — often hours per question.',
        'Ask in plain English. Get a weighted table with counts and percentages in seconds.'
    ),
    (
        'Clients depend entirely on analysts',
        'Clients cannot access SPSS data directly. Every ad-hoc question — even "how many respondents '
        'were aged 25–34?" — requires analyst intervention, creating a bottleneck.',
        'Clients query the data themselves through a browser chat interface, freeing analysts '
        'for higher-value interpretation work.'
    ),
    (
        'Static, one-time reports',
        'Deliverables are fixed PDFs or PowerPoint decks. If a client wants to re-slice data by a '
        'different market or demographic, a new report must be produced — adding days and cost.',
        'Any filter, subgroup, or question can be queried on demand at any time, with no '
        'additional analyst effort.'
    ),
    (
        'Weighting errors and inconsistency',
        'Applying rim weights manually in Excel is error-prone. Analysts may apply weights '
        'inconsistently across questions, leading to discrepancies in delivered figures.',
        'Weights are applied automatically and consistently to every query from the same dataset, '
        'eliminating human error.'
    ),
    (
        'Routed question base confusion',
        'Survey logic means some questions are only shown to a subset of respondents. '
        'Incorrectly using the total sample as the base inflates or deflates percentages.',
        'DataLens automatically detects and reports dual bases — total sample and answered base — '
        'ensuring figures are always statistically correct.'
    ),
    (
        'No self-service trend analysis',
        'Wave-over-wave comparisons require pulling data from multiple datasets, aligning variables, '
        'and computing differences — a multi-hour task for each tracker study.',
        'Ask "How has satisfaction trended across waves?" and receive a wave-by-wave line chart '
        'and table instantly.'
    ),
    (
        'High cost of ad-hoc queries',
        'Every additional question from a client after project delivery represents unbilled analyst '
        'time or a costly scope change. This discourages clients from exploring their data fully.',
        'Unlimited queries at no additional cost per question — clients explore more, get more '
        'value, and come back for the next study.'
    ),
    (
        'Language and variable name barriers',
        'Clients must know SPSS variable names (e.g., "Q11_Lr1") to request specific data cuts. '
        'Most clients have no idea how their data is structured internally.',
        'The platform maps variable codes to plain-language question labels automatically. '
        'Clients ask about "satisfaction with the product" — not "Q11_Lr1".'
    ),
]

for challenge, problem, solution in challenges:
    tbl_p = doc.add_table(rows=1, cols=3)
    tbl_p.style = 'Table Grid'
    widths = [Cm(4.5), Cm(7.5), Cm(7.5)]
    for i, w in enumerate(widths):
        for cell in tbl_p.columns[i].cells:
            cell.width = w

    hdr_row = tbl_p.rows[0]
    # Challenge name cell
    set_cell_bg(hdr_row.cells[0], '1a1f3a')
    cell_text(hdr_row.cells[0], challenge, bold=True, color=WHITE, size=9)
    # Problem cell
    set_cell_bg(hdr_row.cells[1], 'FFF3F3')
    p_cell = hdr_row.cells[1]
    p_cell.text = ''
    pp = p_cell.paragraphs[0]
    r1 = pp.add_run('⚠  Challenge:  ')
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = RED
    r2 = pp.add_run(problem)
    r2.font.size = Pt(8.5)
    # Solution cell
    set_cell_bg(hdr_row.cells[2], 'F0FFF4')
    s_cell = hdr_row.cells[2]
    s_cell.text = ''
    sp = s_cell.paragraphs[0]
    r3 = sp.add_run('✓  DataLens:  ')
    r3.bold = True
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = GREEN
    r4 = sp.add_run(solution)
    r4.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

doc.add_paragraph()
body(
    'In summary: DataLens transforms survey data from a static deliverable into a live, '
    'queryable asset — giving both researchers and clients instant, accurate, self-service '
    'access to insights at a fraction of the current time and cost.',
    bold=True, size=10.5, color=TEAL
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# 3. COMPETITIVE LANDSCAPE
# ═══════════════════════════════════════════════════════════════════════════
heading('3. Competitive Landscape — Module 1 Equivalent Tools', level=1)
divider()
body(
    'The table below covers tools that offer the closest overlap to Module 1\'s core capability: '
    'querying survey data with weighted statistics through an AI or assisted interface. '
    'No single tool currently matches the full combination of features that DataLens delivers.',
    size=10.5
)
doc.add_paragraph()

comp_headers = ['Tool', 'AI / NL Chat', 'SPSS Native', 'Weighted\nStats', 'Modules\nM1/M2/M3', 'Pricing (USD)', 'Key Gap vs DataLens']
comp_rows = [
    # Tool | NL | SPSS | Weighted | Modules | Pricing | Gap
    ['Crunch.io\n(YouGov)',          'No',          'Yes',              'Yes', 'No/No/Yes',  '$20K-$60K/yr',         'Strong crosstabs; zero AI/NL; enterprise-only pricing'],
    ['Infotools Harmoni\n(+Dapresy)','No',          'Yes',              'Yes', 'No/No/Yes',  '$30K-$80K/yr',         'Best for large trackers; no AI; heavy implementation overhead'],
    ['Askia (Cros / Ace)',           'No',          'Yes',              'Yes', 'No/No/Yes',  '$5K-$20K/yr',          'Full survey pipeline; dated UI; no AI or narrative reporting'],
    ['Displayr',                     'No',          'Yes',              'Yes', 'No/No/Yes',  '$1,500-$2,500/user/yr','Widely used in agencies; users still configure tables manually'],
    ['Q Research Software',          'No',          'Yes (best-class)', 'Yes', 'No/No/Yes',  '$1,800-$3,500/user/yr','Gold standard for SPSS stats; syntax-heavy; no client self-service'],
    ['Qualtrics (SAP)',              'Partial',     'Own data only',    'Yes', 'Ptl/No/Ptl', '$1,500-$5,000/user/yr','NL only on Qualtrics-collected data; cannot query external SPSS'],
    ['Forsta / Confirmit',           'Minimal',     'Partial',          'Yes', 'No/No/Yes',  '$50K-$200K/yr',        'End-to-end survey platform; overkill for analysis; very expensive'],
    ['IBM Cognos Analytics',         'Partial (BI)','Partial',          'No',  'Ptl/Ptl/No', '$15-$35/user/month',   'NL for BI dashboards only; not survey-weighted-aware; generic BI'],
    ['IBM SPSS Statistics',          'No',          'Yes',              'Yes', 'No/No/No',   '$1,170/user/yr',       'Definitive SPSS tool; syntax-only; no chatbot; not for clients'],
    ['Yabble',                       'Basic (AI)',  'Partial',          'No',  'Ptl/Ptl/No', '$500-$1,500/month',    'AI summaries for open-ends; no weighted frequency tables or crosstabs'],
    ['ChatGPT / Claude (manual)',    'Yes',         'No metadata',      'No',  'Ptl/No/No',  '$20-$30/user/month',   'Manual upload per session; no SPSS labels; no weighting; no history'],
    ['DataLens (Borderless Access)', 'Full NL',     'Full (labels+codes)','Yes','Yes/Yes/Yes','Custom build\n(Sec 4)','Purpose-built for MR agency SPSS workflows — all 3 modules'],
]

tbl = doc.add_table(rows=1, cols=7)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

col_widths = [Cm(2.8), Cm(1.9), Cm(2.0), Cm(1.7), Cm(1.9), Cm(2.5), Cm(5.0)]
for i, width in enumerate(col_widths):
    for cell in tbl.columns[i].cells:
        cell.width = width

hdr = tbl.rows[0]
for i, h in enumerate(comp_headers):
    cell = hdr.cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, row_data in enumerate(comp_rows):
    row = tbl.add_row()
    bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
    is_datalens = ri == len(comp_rows) - 1
    for ci, val in enumerate(row_data):
        cell = row.cells[ci]
        if is_datalens:
            set_cell_bg(cell, '1a1f3a')
            color = WHITE
        else:
            set_cell_bg(cell, bg)
            color = NAVY if ci == 0 else None
        bold = (ci == 0) or is_datalens
        cell_text(cell, val, bold=bold, color=color, size=8.5,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci in (1,2,3,4) else WD_ALIGN_PARAGRAPH.LEFT)

doc.add_paragraph()
body('M1 = AI NL Chatbot   M2 = Automated Report   M3 = Cross-tab Builder   Ptl = Partial coverage',
     size=8.5, color=GREY, italic=True)
doc.add_paragraph()
body(
    'Key insight: No competing tool offers a true AI chatbot for weighted SPSS survey data. '
    'DataLens Module 1 fills a genuine white space. Combined with Module 2 (AI narrative reports) '
    'and Module 3 (cross-tab builder), DataLens delivers all three capability areas that today '
    'require separate enterprise tools costing $30K to $200K per year.',
    bold=True, size=10.5, color=TEAL
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# 3. INVESTMENT OPTIONS — MODULE 1
# ═══════════════════════════════════════════════════════════════════════════
heading('4. Investment Options — Module 1 Delivery', level=1)
divider()
body(
    'Two delivery approaches are presented below. Both result in a fully functional Module 1 '
    'platform. The right choice depends on the speed required and the availability of internal '
    'analyst capacity.',
    size=10.5
)
doc.add_paragraph()

# ── Option A ────────────────────────────────────────────────────────────────
heading('Option A — With a Dedicated Development Team', level=2)
body(
    'A small specialist team builds and delivers Module 1 to a production-ready state. '
    'This is the faster route and is recommended if a client deadline or internal launch date is fixed.',
    size=10.5
)
doc.add_paragraph()

body('Team Composition', bold=True, size=10.5)
team_a = [
    ('1 × Senior Python / FastAPI Developer', '₹80,000–₹1,20,000 / month', 'Backend, stats engine, API, database'),
    ('1 × React / TypeScript Frontend Developer', '₹70,000–₹1,00,000 / month', 'UI, charts, responsive design, accessibility'),
    ('1 × BA Product Owner (Karthik / team)', 'Internal — no added cost', 'Requirements, domain logic, UAT, sign-off'),
]
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
for i, h in enumerate(['Role', 'Cost (India Market Rate)', 'Responsibility']):
    cell = tbl2.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for role, cost, resp in team_a:
    row = tbl2.add_row()
    for ci, val in enumerate([role, cost, resp]):
        set_cell_bg(row.cells[ci], 'F0F4FF')
        cell_text(row.cells[ci], val, size=9)

doc.add_paragraph()
body('Timeline — Option A', bold=True, size=10.5)

timeline_a = [
    ('Sprint 1', '2 weeks', 'Auth & access controls (JWT login, role-based access, project expiry)'),
    ('Sprint 2', '2 weeks', 'Responsive UI (mobile / tablet breakpoints, collapsible sidebar)'),
    ('Sprint 3', '2 weeks', 'Performance, error handling, data validation, multi-file uploads'),
    ('Sprint 4', '2 weeks', 'UAT, bug fixes, staging deployment, client demo'),
    ('Sprint 5', '2 weeks', 'Production deployment, monitoring, handover documentation'),
    ('Buffer', '1 week', 'Contingency'),
]
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
for i, h in enumerate(['Sprint', 'Duration', 'Deliverables']):
    cell = tbl3.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for sprint, dur, work in timeline_a:
    row = tbl3.add_row()
    for ci, val in enumerate([sprint, dur, work]):
        set_cell_bg(row.cells[ci], 'F0F4FF')
        cell_text(row.cells[ci], val, size=9, bold=(ci==0))

doc.add_paragraph()
body('Cost Summary — Option A', bold=True, size=10.5)

cost_a = [
    ('Development team (2 devs × ~2.5 months)', '₹3,75,000 – ₹5,50,000', '~$4,500 – $6,600'),
    ('Anthropic API usage (Claude Sonnet)', '₹4,200 – ₹12,500 / month', '~$50 – $150/month'),
    ('Cloud hosting (Azure / AWS)', '₹8,400 – ₹25,000 / month', '~$100 – $300/month'),
    ('Database (SQL Server / PostgreSQL)', '₹4,200 – ₹12,500 / month', '~$50 – $150/month'),
    ('Total one-time build cost', '₹3,75,000 – ₹5,50,000', '~$4,500 – $6,600'),
    ('Total monthly running cost', '₹16,800 – ₹50,000 / month', '~$200 – $600/month'),
]
tbl4 = doc.add_table(rows=1, cols=3)
tbl4.style = 'Table Grid'
for i, h in enumerate(['Cost Item', 'INR', 'USD']):
    cell = tbl4.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for ri, (item, inr, usd) in enumerate(cost_a):
    row = tbl4.add_row()
    is_total = item.startswith('Total')
    bg = '1a1f3a' if is_total else ('F0F4FF' if ri % 2 == 0 else 'FFFFFF')
    col = WHITE if is_total else None
    for ci, val in enumerate([item, inr, usd]):
        set_cell_bg(row.cells[ci], bg)
        cell_text(row.cells[ci], val, bold=is_total, color=col, size=9)

doc.add_paragraph()
body('Estimated delivery: 10–11 weeks from kick-off to production.', bold=True, size=10.5, color=TEAL)

doc.add_paragraph()
divider()

# ── Option B ────────────────────────────────────────────────────────────────
heading('Option B — Data Analysts + Claude Code (No Development Team)', level=2)
body(
    'Internal data analysts at Borderless Access use Claude Code (Anthropic\'s AI coding assistant) '
    'to build and iterate on Module 1 without hiring external developers. This approach has already '
    'delivered the current working POC and is suitable for teams comfortable with guided AI-assisted '
    'development. It has a lower cash investment but requires more analyst time and a longer timeline.',
    size=10.5
)
doc.add_paragraph()

body('Team Composition', bold=True, size=10.5)
team_b = [
    ('1–2 × Data Analysts (BA internal)', 'Internal — existing headcount', '~30–50% of working time on DataLens build'),
    ('Claude Code (Anthropic Max plan)', '$100 / analyst / month', 'AI pair programmer — writes, tests, debugs code'),
    ('Anthropic API (platform usage)', '$50–$150 / month', 'Powers the DataLens chatbot queries'),
    ('Cloud hosting', '$100–$300 / month', 'Azure / AWS for backend + frontend'),
]
tbl5 = doc.add_table(rows=1, cols=3)
tbl5.style = 'Table Grid'
for i, h in enumerate(['Role / Resource', 'Cost', 'Notes']):
    cell = tbl5.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for role, cost, note in team_b:
    row = tbl5.add_row()
    for ci, val in enumerate([role, cost, note]):
        set_cell_bg(row.cells[ci], 'F0F4FF')
        cell_text(row.cells[ci], val, size=9)

doc.add_paragraph()
body('Timeline — Option B', bold=True, size=10.5)

timeline_b = [
    ('Month 1', 'Auth & access controls, bug fixes from POC testing'),
    ('Month 2', 'Responsive design, UI polish, performance improvements'),
    ('Month 3', 'UAT, additional question-type support, client demo preparation'),
    ('Month 4', 'Staging deployment, documentation, training'),
    ('Month 5', 'Production launch, monitoring, iterative improvements'),
]
tbl6 = doc.add_table(rows=1, cols=2)
tbl6.style = 'Table Grid'
for i, h in enumerate(['Period', 'Focus']):
    cell = tbl6.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for period, focus in timeline_b:
    row = tbl6.add_row()
    for ci, val in enumerate([period, focus]):
        set_cell_bg(row.cells[ci], 'F0F4FF')
        cell_text(row.cells[ci], val, size=9, bold=(ci==0))

doc.add_paragraph()
body('Cost Summary — Option B', bold=True, size=10.5)

cost_b = [
    ('Claude Code license (Max plan, 1–2 analysts)', '$100–$200 / month', 'During active build phase'),
    ('Anthropic API (platform usage)', '$50–$150 / month', 'Ongoing once live'),
    ('Cloud hosting', '$100–$300 / month', 'Ongoing'),
    ('Developer salaries', '$0 additional', 'Uses existing analyst headcount'),
    ('Total monthly cost (build phase)', '$250–$650 / month', '~₹21,000 – ₹54,000'),
    ('Total monthly cost (post-launch)', '$150–$450 / month', '~₹12,500 – ₹37,500 (no Claude Code needed)'),
]
tbl7 = doc.add_table(rows=1, cols=3)
tbl7.style = 'Table Grid'
for i, h in enumerate(['Cost Item', 'USD', 'Notes']):
    cell = tbl7.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for ri, (item, usd, note) in enumerate(cost_b):
    row = tbl7.add_row()
    is_total = item.startswith('Total')
    bg = '1a1f3a' if is_total else ('F0F4FF' if ri % 2 == 0 else 'FFFFFF')
    col = WHITE if is_total else None
    for ci, val in enumerate([item, usd, note]):
        set_cell_bg(row.cells[ci], bg)
        cell_text(row.cells[ci], val, bold=is_total, color=col, size=9)

doc.add_paragraph()
body('Estimated delivery: 4–5 months from start to production.', bold=True, size=10.5, color=TEAL)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# 4. OPTION COMPARISON
# ═══════════════════════════════════════════════════════════════════════════
heading('5. Option Comparison at a Glance', level=1)
divider()
doc.add_paragraph()

compare = [
    ('Delivery timeline',        '10–11 weeks',                        '4–5 months'),
    ('One-time build cost',      '₹3.75L – ₹5.5L  (~$4,500–$6,600)',   'Near zero (uses existing headcount)'),
    ('Monthly running cost',     '$200–$600 / month',                   '$150–$450 / month (post-launch)'),
    ('Technical risk',           'Low — dedicated specialists',         'Medium — analysts learning as they build'),
    ('Internal capability built','Low — external team delivers',        'High — analysts gain AI dev skills'),
    ('Speed to client demo',     'Fast — ~6 weeks',                    'Slower — ~10 weeks'),
    ('Recommended for',          'Fixed client deadlines / revenue pressure', 'Internal tooling / capability building'),
]

tbl8 = doc.add_table(rows=1, cols=3)
tbl8.style = 'Table Grid'
for i, h in enumerate(['Factor', 'Option A: Dev Team', 'Option B: Analysts + Claude Code']):
    cell = tbl8.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for ri, (factor, a, b) in enumerate(compare):
    row = tbl8.add_row()
    bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
    for ci, val in enumerate([factor, a, b]):
        set_cell_bg(row.cells[ci], bg)
        cell_text(row.cells[ci], val, size=9, bold=(ci==0))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# 5. FUTURE ROADMAP — MODULE 2 & 3
# ═══════════════════════════════════════════════════════════════════════════
heading('6. Future Roadmap — Module 2 & Module 3', level=1)
divider()
body(
    'Module 1 lays the data and infrastructure foundation. Modules 2 and 3 build on the same '
    'backend and extend the platform into automated reporting and self-service analysis.',
    size=10.5
)
doc.add_paragraph()

# Module 2
heading('Module 2 — Automated Summary Report Generator', level=2)
body(
    'Generates a ready-to-share insight deck from a survey dataset with one click. '
    'Claude reads the full question registry, runs key statistics, interprets results, '
    'and outputs a structured PowerPoint or PDF report with narrative commentary and charts.',
    size=10.5
)
doc.add_paragraph()
body('Key features:', bold=True, size=10.5)
bullet('Auto-select significant findings per question (top responses, notable shifts)')
bullet('Section-by-section narrative: demographics → attitudes → behaviours → recommendations')
bullet('Branded output — Borderless Access or client logo, colour scheme')
bullet('Editable PPTX export — researchers can refine before sharing')
bullet('Optional: compare two waves or two sub-groups side-by-side')

doc.add_paragraph()

mod2_plan = [
    ('Estimated timeline', 'Option A (dev team): 6–8 weeks | Option B (analysts): 3–4 months'),
    ('Estimated build cost (dev team)', '₹2.5L – ₹4L  (~$3,000–$4,800)'),
    ('Additional monthly API cost', '+$100–$300/month (higher token usage for full-report generation)'),
    ('Dependencies', 'Module 1 fully deployed; python-pptx or similar report library'),
]
tbl9 = doc.add_table(rows=len(mod2_plan), cols=2)
tbl9.style = 'Table Grid'
for ri, (k, v) in enumerate(mod2_plan):
    bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
    set_cell_bg(tbl9.rows[ri].cells[0], bg)
    set_cell_bg(tbl9.rows[ri].cells[1], bg)
    cell_text(tbl9.rows[ri].cells[0], k, bold=True, size=9)
    cell_text(tbl9.rows[ri].cells[1], v, size=9)

doc.add_paragraph()

# Module 3
heading('Module 3 — Self-Service Cross-Tab Builder', level=2)
body(
    'A drag-and-drop interface that allows researchers and clients to build custom cross-tabulations '
    'without writing any code or syntax. Users select row variables, column variables, filters, and '
    'significance tests — and the platform computes and displays the table instantly.',
    size=10.5
)
doc.add_paragraph()
body('Key features:', bold=True, size=10.5)
bullet('Drag-and-drop variable selector for rows (banner) and columns (stub)')
bullet('Weighted counts and column percentages with significance indicators (arrows or asterisks)')
bullet('Filters by any variable (age, gender, region, usage, etc.)')
bullet('Export to Excel (.xlsx) with formatting preserved')
bullet('Save cross-tab configurations and re-run on updated data')
bullet('Optional: net difference column, index vs total')

doc.add_paragraph()

mod3_plan = [
    ('Estimated timeline', 'Option A (dev team): 8–10 weeks | Option B (analysts): 4–6 months'),
    ('Estimated build cost (dev team)', '₹4L – ₹6.5L  (~$4,800–$7,800)  — heavier frontend work'),
    ('Additional monthly API cost', 'Minimal — crosstabs are computed locally, not via Claude'),
    ('Dependencies', 'Module 1 deployed; React drag-drop library (dnd-kit); stats engine extension'),
]
tbl10 = doc.add_table(rows=len(mod3_plan), cols=2)
tbl10.style = 'Table Grid'
for ri, (k, v) in enumerate(mod3_plan):
    bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
    set_cell_bg(tbl10.rows[ri].cells[0], bg)
    set_cell_bg(tbl10.rows[ri].cells[1], bg)
    cell_text(tbl10.rows[ri].cells[0], k, bold=True, size=9)
    cell_text(tbl10.rows[ri].cells[1], v, size=9)

doc.add_paragraph()

# Full roadmap timeline
heading('Full Platform Timeline (All 3 Modules)', level=2)
roadmap = [
    ('Now',          'Module 1 POC — functional, internal testing',          '✓ Complete'),
    ('Month 1–3',    'Module 1 production: auth, hosting, responsive UI',    'In Progress'),
    ('Month 4–5',    'Module 2: Automated report generator',                 'Planned'),
    ('Month 6–8',    'Module 3: Cross-tab builder',                          'Planned'),
    ('Month 9+',     'Phase B: Multi-tenant, client portal, billing, API',   'Future'),
]
tbl11 = doc.add_table(rows=1, cols=3)
tbl11.style = 'Table Grid'
for i, h in enumerate(['Timeline', 'Milestone', 'Status']):
    cell = tbl11.rows[0].cells[i]
    set_cell_bg(cell, '1a1f3a')
    cell_text(cell, h, bold=True, color=WHITE, size=9)
for ri, (period, milestone, status) in enumerate(roadmap):
    row = tbl11.add_row()
    bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
    status_color = GREEN if '✓' in status else (TEAL if 'Progress' in status else AMBER)
    for ci, val in enumerate([period, milestone, status]):
        set_cell_bg(row.cells[ci], bg)
        col = status_color if ci == 2 else None
        cell_text(row.cells[ci], val, size=9, bold=(ci in (0, 2)), color=col)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════
# 6. RECOMMENDATION
# ═══════════════════════════════════════════════════════════════════════════
heading('7. Recommendation', level=1)
divider()
doc.add_paragraph()
body(
    'For Borderless Access, Option B (Data Analysts + Claude Code) is recommended as the primary '
    'path for Module 1. The POC has already demonstrated this approach works. The marginal '
    'additional investment to reach production is low, and the team builds lasting AI-assisted '
    'development capability that benefits all three modules.',
    size=11
)
doc.add_paragraph()
body(
    'If a specific client or revenue opportunity requires a faster launch (within 10 weeks), '
    'Option A (Dedicated Dev Team) should be activated alongside the internal team.',
    size=11
)
doc.add_paragraph()
body('Key advantages of building over buying:', bold=True, size=10.5)
bullet('No per-user annual licence fees — marginal cost per researcher is near zero')
bullet('Full control over SPSS metadata, weighting logic, and BA-specific conventions')
bullet('Data stays inside BA infrastructure — no third-party data sharing')
bullet('Extensible — Module 2 and 3 build on the same platform at low incremental cost')
bullet('Potential to offer DataLens to clients as a differentiated service')

doc.add_paragraph()
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_final.add_run('Powered by Borderless Access  |  DataLens AI Survey Insights Platform')
r.font.size = Pt(9)
r.font.color.rgb = GREY
r.font.italic = True

doc.save(OUT)
print('Saved:', OUT)
